"""Сборка DOCX из результата диаризации (python-docx).

Вход — словарь как в JSON-ответе: {summary?, follow_up?, todo?, segments, timestamps}.
"""
from __future__ import annotations

import io
import re

from docx import Document


def _mmss(sec) -> str:
    s = int(sec or 0)
    return f"{s // 60:02d}:{s % 60:02d}"


def _add_runs(p, text: str) -> None:
    """Текст с **жирным** → runs."""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text or "")):
        if part:
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True


def _add_md(doc, text: str, numbered: bool = False) -> None:
    """Markdown-ish: пункты `- ...` → списком, остальные непустые строки → абзацами."""
    style = "List Number" if numbered else "List Bullet"
    for line in (text or "").split("\n"):
        m = re.match(r"^(\s*)[-•*]\s+(.*)$", line)
        if m:
            _add_runs(doc.add_paragraph(style=style), m.group(2))
        elif line.strip():
            _add_runs(doc.add_paragraph(), line.strip())


def _parse_todo(text: str):
    rows = []
    for line in (text or "").split("\n"):
        line = line.strip()
        if "|" not in line:
            continue
        cells = [c.strip() for c in re.sub(r"^\||\|$", "", line).split("|")]
        if all(re.match(r"^[-:\s]*$", c) for c in cells):
            continue  # строка-разделитель |---|
        rows.append(cells)
    if len(rows) > 1 and re.search(r"задач|ответствен|срок", " ".join(rows[0]), re.I):
        rows = rows[1:]
    return rows


def _add_todo(doc, text: str) -> None:
    rows = _parse_todo(text)
    if not rows:
        _add_md(doc, text)
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["№", "Задача", "Ответственный", "Срок"]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
    for i, c in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = str(i + 1)
        _add_runs(cells[1].paragraphs[0], c[0] if len(c) > 0 else "")
        _add_runs(cells[2].paragraphs[0], c[1] if len(c) > 1 else "—")
        _add_runs(cells[3].paragraphs[0], c[2] if len(c) > 2 else "—")


def build_docx(data: dict) -> bytes:
    doc = Document()
    if data.get("summary"):
        doc.add_heading("Саммари", level=1)
        _add_md(doc, data["summary"])
    if data.get("follow_up"):
        doc.add_heading("Follow-up", level=1)
        _add_md(doc, data["follow_up"], numbered=True)
    if data.get("todo"):
        doc.add_heading("To-do", level=1)
        _add_todo(doc, data["todo"])
    doc.add_heading("Транскрипт", level=1)
    ts = data.get("timestamps", True)
    for u in data.get("segments", []):
        p = doc.add_paragraph()
        prefix = f"[{_mmss(u.get('start'))}-{_mmss(u.get('end'))}] " if ts else ""
        p.add_run(prefix + str(u.get("speaker", "")) + ": ").bold = True
        p.add_run(str(u.get("text", "")))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
